"""
ui/exporter.py

1-Click Export module for the Multimodal AI Media Summarizer.
Exports AI-generated summaries in three formats:
  - PDF via ReportLab (emoji-safe, professional layout, zero system deps)
  - DOCX via python-docx (preserves full heading hierarchy, nested bullets, blockquotes, bold/italic)
  - Markdown via direct UTF-8 encoding (zero-loss, reference format)

Session state keys used (read-only):
  st.session_state.result.summary_markdown: str — raw markdown summary
  st.session_state.result.video_filename:   str — original source filename

All exports are generated in-memory (BytesIO) — zero disk writes.
Zero-footprint: no temp files created or left behind.
"""

from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO
from typing import Any

import streamlit as st

# ── ReportLab (PDF) ──────────────────────────────────────────────────────────
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

# ── python-docx (DOCX) ───────────────────────────────────────────────────────
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ═══════════════════════════════════════════════════════════════════════════════
# EMOJI MAP — only used for PDF (ReportLab cannot render emoji glyphs)
# DOCX and Markdown pass emoji through unchanged.
# ═══════════════════════════════════════════════════════════════════════════════
_EMOJI_MAP: dict[str, str] = {
    # Prompt schema headers
    "🎬": "[MEDIA]",
    "📋": "[INFO]",
    "🗂️": "[TOC]",
    "📝": "[NOTES]",
    "🔑": "[KEY]",
    "📚": "[DETAIL]",
    "💡": "[INSIGHT]",
    "❓": "[Q&A]",
    "🌐": "[GLOBAL]",
    "🏷️": "[TAGS]",
    # Lecture schema headers (user-specified)
    "🎓": "[OVERVIEW]",
    "📌": "[CONCEPTS]",
    "📖": "[BREAKDOWN]",
    "⚠️": "[WARNING]",
    "🔗": "[MAP]",
    # Other common emoji
    "✅": "[OK]",
    "❌": "[FAIL]",
    "→": "->",
    "•": "*",
    "–": "-",
    "—": "-",
    "…": "...",
}


def _clean_math(text: str) -> str:
    """
    Clean LaTeX math notation and escaped symbols so text/DOCX/PDF exports
    render readable mathematical and set notation instead of raw TeX backslashes.
    E.g. $\\{A, B\\}$ -> {A, B}, \\rightarrow -> →, \\dots -> …
    """
    if not text:
        return ""

    # Common LaTeX symbol replacements
    replacements = [
        (r"\\\{", "{"),
        (r"\\\}", "}"),
        (r"\\rightarrow", "→"),
        (r"\\to\b", "→"),
        (r"\\leftarrow", "←"),
        (r"\\Rightarrow", "⇒"),
        (r"\\Leftarrow", "⇐"),
        (r"\\Leftrightarrow", "⇔"),
        (r"\\in\b", "∈"),
        (r"\\notin\b", "∉"),
        (r"\\subset\b", "⊂"),
        (r"\\subseteq\b", "⊆"),
        (r"\\cup\b", "∪"),
        (r"\\cap\b", "∩"),
        (r"\\emptyset\b", "∅"),
        (r"\\dots\b", "…"),
        (r"\\ldots\b", "…"),
        (r"\\approx\b", "≈"),
        (r"\\neq\b", "≠"),
        (r"\\leq\b", "≤"),
        (r"\\le\b", "≤"),
        (r"\\geq\b", "≥"),
        (r"\\ge\b", "≥"),
        (r"\\\$", "$"),
        (r"\\\_", "_"),
    ]
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text)

    # Strip inline math $...$ delimiters for set/variable/math expressions
    # e.g., ${A, B}$ -> {A, B}, $B$ -> B, $x = 5$ -> x = 5
    # Avoid stripping currency like $100 or $50.00
    def _unwrap_math(m: re.Match[str]) -> str:
        content = m.group(1)
        if re.match(r"^\d+(?:\.\d+)?$", content):
            return m.group(0)
        return content

    text = re.sub(r"\$([^\$\n]+)\$", _unwrap_math, text)
    return text


def _clean_for_pdf(text: str) -> str:
    """
    Replace known emoji with text equivalents, then strip any remaining
    non-Latin-extended characters that ReportLab cannot render.

    Args:
        text: Raw text that may contain emoji.

    Returns:
        PDF-safe string with emoji replaced or removed.
    """
    for emoji, replacement in _EMOJI_MAP.items():
        text = text.replace(emoji, replacement)

    # Strip any remaining codepoints in emoji/symbol ranges
    result: list[str] = []
    for char in text:
        cp = ord(char)
        # Allow: Basic Latin, Latin Extended, common punctuation/math/arrows
        if cp < 0x2000 or (0x2190 <= cp <= 0x21FF) or (0x2200 <= cp <= 0x22FF):
            result.append(char)
        elif char in ("✓", "°", "©", "®", "™"):
            result.append(char)
        else:
            result.append("")   # silently drop unknown glyph
    return "".join(result)


def _inline_pdf(text: str) -> str:
    """
    Convert markdown inline markers to ReportLab XML tags.
    Order: bold-italic before bold/italic (greedy prevention).

    Args:
        text: Markdown text with **bold**, *italic*, `code`.

    Returns:
        ReportLab Paragraph XML string.
    """
    # Escape reserved XML chars first
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # Bold-italic: ***text***
    text = re.sub(r"\*\*\*(.*?)\*\*\*", r"<b><i>\1</i></b>", text)
    # Bold: **text**
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    # Italic: *text*
    text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
    # Code: `text`
    text = re.sub(r"`(.*?)`", r'<font name="Courier" size="9">\1</font>', text)
    return text


def _add_formatted_runs(para: Any, text: str) -> None:
    """
    Parse inline markdown **bold**, *italic*, `code`, ***bold-italic***
    and add correctly formatted runs to a python-docx Paragraph.

    Args:
        para: python-docx Paragraph object.
        text: Markdown text to parse and add.
    """
    pattern = r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)"
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***") and len(part) > 6:
            run = para.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            para.add_run(part)


def _add_paragraph_border(para: Any) -> None:
    """
    Add a blue left border to a python-docx paragraph (blockquote style).

    Args:
        para: python-docx Paragraph object.
    """
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), "1E40AF")
    pBdr.append(left)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT FUNCTION 1 — MARKDOWN
# ═══════════════════════════════════════════════════════════════════════════════

def export_markdown(summary_text: str, filename: str) -> bytes:
    """
    Encode summary text as UTF-8 Markdown bytes with a YAML front-matter header.

    Cleans raw LaTeX math escapes (e.g. \\{A, B\\} -> {A, B}) while preserving
    full markdown formatting and structure.

    Args:
        summary_text: Raw markdown string from the AI.
        filename: Base filename without extension.

    Returns:
        UTF-8 encoded bytes ready for st.download_button().
    """
    cleaned = _clean_math(summary_text)
    header = (
        f"---\n"
        f"title: AI Media Summary\n"
        f"generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"format: Markdown\n"
        f"---\n\n"
    )
    return (header + cleaned).encode("utf-8")


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT FUNCTION 2 — DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def export_docx(summary_text: str, filename: str) -> bytes:
    """
    Convert markdown summary to a formatted Microsoft Word (.docx) document.

    Preserves full heading hierarchy (H1–H4), nested bullet/numbered lists,
    blockquotes (with a blue left border), inline bold/italic/code, and
    horizontal rules. Cleans LaTeX math notation. Emoji pass through natively.

    Args:
        summary_text: Raw markdown string from the AI.
        filename: Base filename without extension.

    Returns:
        DOCX bytes (ZIP/PK magic bytes) ready for st.download_button().
    """
    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Document title block ────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("AI Media Summary")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}"
    )
    sub_run.font.size = Pt(10)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph("─" * 60)

    # ── Parse markdown line by line ─────────────────────────────────────────
    lines = summary_text.split("\n")
    i = 0
    while i < len(lines):
        raw_line = lines[i].rstrip()
        line = _clean_math(raw_line)
        stripped = line.strip()

        if stripped.startswith("#### "):
            h = doc.add_heading(level=4)
            _add_formatted_runs(h, stripped[5:].strip())

        elif stripped.startswith("### "):
            h = doc.add_heading(level=3)
            r = h.add_run(stripped[4:].strip())
            r.font.color.rgb = RGBColor(0x37, 0x4B, 0xA0)

        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            r = h.add_run(stripped[3:].strip())
            r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

        elif stripped.startswith("# "):
            h = doc.add_heading(level=1)
            r = h.add_run(stripped[2:].strip())
            r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

        elif match := re.match(r"^(\s*)([-*•])\s+(.*)", line):
            indent_str, _, text = match.groups()
            indent_spaces = len(indent_str.expandtabs(4))
            para = doc.add_paragraph(style="List Bullet")
            if indent_spaces > 0:
                indent_level = max(1, indent_spaces // 2)
                para.paragraph_format.left_indent = Inches(0.25 * indent_level + 0.25)
            _add_formatted_runs(para, text)

        elif match := re.match(r"^(\s*)(\d+)[\.\)]\s+(.*)", line):
            indent_str, num_str, text = match.groups()
            indent_spaces = len(indent_str.expandtabs(4))
            para = doc.add_paragraph(style="List Number")
            if indent_spaces > 0:
                indent_level = max(1, indent_spaces // 2)
                para.paragraph_format.left_indent = Inches(0.25 * indent_level + 0.25)
            _add_formatted_runs(para, text)

        elif match := re.match(r"^(\s*)>\s*(.*)", line):
            _, text = match.groups()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            _add_formatted_runs(para, text)
            _add_paragraph_border(para)
            for run in para.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x37, 0x4B, 0xA0)

        elif stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)

        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)

        elif stripped == "":
            pass  # skip empty lines — no extra blank paragraphs

        else:
            if stripped:
                para = doc.add_paragraph()
                _add_formatted_runs(para, line)

        i += 1

    # ── Footer ──────────────────────────────────────────────────────────────
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f"Generated by Multimodal AI Media Summarizer  ·  "
        f"{datetime.now().strftime('%Y-%m-%d')}"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT FUNCTION 3 — PDF
# ═══════════════════════════════════════════════════════════════════════════════

def export_pdf(summary_text: str, filename: str) -> bytes:
    """
    Convert markdown summary to a professional PDF document via ReportLab.

    Emoji are replaced with text equivalents before rendering to prevent
    ReportLab crashes. Preserves heading hierarchy, nested bullet lists,
    blockquotes, inline bold/italic/code, and horizontal rules. Cleans LaTeX math.

    Args:
        summary_text: Raw markdown string from the AI.
        filename: Base filename without extension.

    Returns:
        PDF bytes (%PDF magic bytes) ready for st.download_button().
    """
    buf = BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=0.75 * inch,
        title="AI Media Summary",
        author="Multimodal AI Media Summarizer",
    )

    # ── Custom styles ────────────────────────────────────────────────────────
    styles = getSampleStyleSheet()

    s_title = ParagraphStyle(
        "DocTitle", parent=styles["Normal"],
        fontSize=26, fontName="Helvetica-Bold",
        textColor=HexColor("#1A56DB"),
        alignment=TA_CENTER, spaceAfter=4,
    )
    s_sub = ParagraphStyle(
        "DocSub", parent=styles["Normal"],
        fontSize=10, textColor=HexColor("#6B7280"),
        alignment=TA_CENTER, spaceAfter=16,
    )
    s_h1 = ParagraphStyle(
        "H1", parent=styles["Normal"],
        fontSize=18, fontName="Helvetica-Bold",
        textColor=HexColor("#1A56DB"),
        spaceBefore=18, spaceAfter=8,
    )
    s_h2 = ParagraphStyle(
        "H2", parent=styles["Normal"],
        fontSize=14, fontName="Helvetica-Bold",
        textColor=HexColor("#1E40AF"),
        spaceBefore=14, spaceAfter=6,
    )
    s_h3 = ParagraphStyle(
        "H3", parent=styles["Normal"],
        fontSize=12, fontName="Helvetica-Bold",
        textColor=HexColor("#374BA0"),
        spaceBefore=10, spaceAfter=4,
    )
    s_h4 = ParagraphStyle(
        "H4", parent=styles["Normal"],
        fontSize=11, fontName="Helvetica-Bold",
        textColor=HexColor("#4B5563"),
        spaceBefore=8, spaceAfter=4, leftIndent=12,
    )
    s_body = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, textColor=HexColor("#1F2937"),
        leading=15, spaceAfter=5, alignment=TA_JUSTIFY,
    )
    s_bullet = ParagraphStyle(
        "Bullet", parent=styles["Normal"],
        fontSize=10, textColor=HexColor("#1F2937"),
        leading=14, spaceAfter=3, leftIndent=20,
    )
    s_quote = ParagraphStyle(
        "Quote", parent=styles["Normal"],
        fontSize=10, fontName="Helvetica-Oblique",
        textColor=HexColor("#374BA0"),
        leftIndent=24, rightIndent=12,
        spaceBefore=6, spaceAfter=6, leading=15,
        backColor=HexColor("#EFF6FF"),
        borderColor=HexColor("#1E40AF"),
        borderWidth=3, borderPad=8,
    )
    s_code = ParagraphStyle(
        "Code", parent=styles["Code"],
        fontSize=9, fontName="Courier",
        backColor=HexColor("#F3F4F6"),
        leftIndent=16, rightIndent=8,
        spaceBefore=4, spaceAfter=8, leading=13,
    )

    # ── Page footer callback ─────────────────────────────────────────────────
    def _footer(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(HexColor("#9CA3AF"))
        canvas.drawCentredString(
            letter[0] / 2,
            0.4 * inch,
            f"Multimodal AI Media Summarizer  ·  "
            f"Page {doc.page}  ·  "
            f"Generated {datetime.now().strftime('%Y-%m-%d')}",
        )
        canvas.restoreState()

    # ── Parse markdown → flowables ───────────────────────────────────────────
    story: list[Any] = []

    story.append(Paragraph("AI Media Summary", s_title))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}",
        s_sub,
    ))
    story.append(HRFlowable(
        width="100%", thickness=2,
        color=HexColor("#1A56DB"), spaceAfter=14,
    ))

    lines = summary_text.split("\n")
    i = 0
    bullet_buf: list[tuple[str, int]] = []

    def _flush_bullets() -> None:
        if not bullet_buf:
            return
        items = []
        for text, level in bullet_buf:
            indent_pt = 16 + (level * 12)
            b_style = ParagraphStyle(
                f"Bullet_{indent_pt}",
                parent=s_bullet,
                leftIndent=indent_pt,
            )
            items.append(
                ListItem(
                    Paragraph(_inline_pdf(text), b_style),
                    bulletColor=HexColor("#1A56DB"),
                    bulletType="bullet",
                )
            )
        story.append(ListFlowable(
            items,
            bulletType="bullet",
            start="•",
            leftIndent=16,
            bulletFontSize=10,
            bulletColor=HexColor("#1A56DB"),
        ))
        bullet_buf.clear()

    while i < len(lines):
        raw_line = lines[i].rstrip()
        line = _clean_math(raw_line)
        stripped = line.strip()

        if stripped.startswith("#### "):
            _flush_bullets()
            story.append(Paragraph(_inline_pdf(_clean_for_pdf(stripped[5:].strip())), s_h4))

        elif stripped.startswith("### "):
            _flush_bullets()
            story.append(Spacer(1, 4))
            story.append(Paragraph(_inline_pdf(_clean_for_pdf(stripped[4:].strip())), s_h3))

        elif stripped.startswith("## "):
            _flush_bullets()
            story.append(Spacer(1, 6))
            story.append(Paragraph(_inline_pdf(_clean_for_pdf(stripped[3:].strip())), s_h2))

        elif stripped.startswith("# "):
            _flush_bullets()
            story.append(Spacer(1, 8))
            story.append(Paragraph(_inline_pdf(_clean_for_pdf(stripped[2:].strip())), s_h1))
            story.append(HRFlowable(
                width="100%", thickness=1,
                color=HexColor("#DBEAFE"), spaceAfter=4,
            ))

        elif match := re.match(r"^(\s*)([-*•])\s+(.*)", line):
            indent_str, _, text = match.groups()
            indent_spaces = len(indent_str.expandtabs(4))
            indent_level = max(0, indent_spaces // 2)
            bullet_buf.append((_clean_for_pdf(text), indent_level))

        elif match := re.match(r"^(\s*)(\d+)[\.\)]\s+(.*)", line):
            _flush_bullets()
            indent_str, num_str, text = match.groups()
            indent_spaces = len(indent_str.expandtabs(4))
            indent_pt = 16 + max(0, indent_spaces // 2) * 12
            story.append(Paragraph(
                f"{num_str}. {_inline_pdf(_clean_for_pdf(text))}",
                ParagraphStyle(
                    f"Num_{indent_pt}", parent=s_bullet, leftIndent=indent_pt
                )
            ))

        elif match := re.match(r"^(\s*)>\s*(.*)", line):
            _flush_bullets()
            _, text = match.groups()
            story.append(Paragraph(
                _inline_pdf(_clean_for_pdf(text)), s_quote
            ))

        elif stripped.startswith("```"):
            _flush_bullets()
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                story.append(Paragraph(
                    "<br/>".join(_clean_for_pdf(cl) for cl in code_lines),
                    s_code,
                ))

        elif stripped in ("---", "***", "___"):
            _flush_bullets()
            story.append(HRFlowable(
                width="100%", thickness=1,
                color=HexColor("#E5E7EB"), spaceAfter=6,
            ))

        elif stripped == "":
            _flush_bullets()
            story.append(Spacer(1, 4))

        else:
            if stripped:
                _flush_bullets()
                story.append(Paragraph(
                    _inline_pdf(_clean_for_pdf(line)), s_body
                ))

        i += 1

    _flush_bullets()

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# MASTER RENDER FUNCTION
# ═══════════════════════════════════════════════════════════════════════════════

def render_export_buttons(
    summary_text: str,
    base_filename: str | None = None,
) -> None:
    """
    Render three side-by-side 1-click export buttons below the summary.

    Generates PDF, DOCX, and Markdown exports on demand, entirely in memory.
    Buttons are styled with distinct red/blue/green color accents.

    Args:
        summary_text: Raw markdown summary from the AI.
        base_filename: Optional base filename (without extension).
                       Defaults to a timestamp-based name.
    """
    if not summary_text or not summary_text.strip():
        return  # Nothing to export

    if not base_filename:
        base_filename = f"ai_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    # Sanitize: keep alphanumeric, hyphen, underscore only
    safe_name = re.sub(r"[^\w\-]", "_", base_filename)

    # ── Section header ───────────────────────────────────────────────────────
    st.markdown(
        """
        <div style="
            margin: 28px 0 14px 0;
            display: flex;
            align-items: center;
            gap: 10px;
        ">
            <span style="
                font-size: 11px;
                letter-spacing: 2px;
                text-transform: uppercase;
                color: rgba(255,255,255,0.35);
                font-family: Inter, system-ui, sans-serif;
                font-weight: 600;
            ">📥 Export Summary</span>
            <div style="flex: 1; height: 1px; background: rgba(255,255,255,0.06);"></div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    col_pdf, col_docx, col_md = st.columns(3)

    # ── PDF ──────────────────────────────────────────────────────────────────
    with col_pdf:
        try:
            pdf_bytes = export_pdf(summary_text, safe_name)
            st.download_button(
                label="📄 Download PDF",
                data=pdf_bytes,
                file_name=f"{safe_name}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="export_pdf_btn",
            )
        except Exception as exc:
            st.error(f"PDF generation failed: {exc}")

    # ── DOCX ─────────────────────────────────────────────────────────────────
    with col_docx:
        try:
            docx_bytes = export_docx(summary_text, safe_name)
            st.download_button(
                label="📝 Download Word",
                data=docx_bytes,
                file_name=f"{safe_name}.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"
                ),
                use_container_width=True,
                key="export_docx_btn",
            )
        except Exception as exc:
            st.error(f"DOCX generation failed: {exc}")

    # ── Markdown ─────────────────────────────────────────────────────────────
    with col_md:
        try:
            md_bytes = export_markdown(summary_text, safe_name)
            st.download_button(
                label="⬇️ Download Markdown",
                data=md_bytes,
                file_name=f"{safe_name}.md",
                mime="text/markdown",
                use_container_width=True,
                key="export_md_btn",
            )
        except Exception as exc:
            st.error(f"Markdown export failed: {exc}")
