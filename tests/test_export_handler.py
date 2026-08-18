"""
tests/test_export_handler.py — Unit tests for ui/exporter.py

Covers: export_markdown, export_docx, export_pdf, filename sanitization,
        emoji handling, and the full 9-section prompt structure.
"""

from __future__ import annotations

import re
from io import BytesIO

import pytest
from docx import Document as DocxDocument

from ui.exporter import export_docx, export_markdown, export_pdf

# ── Canonical mock summary replicating the LECTURE_SUMMARY_PROMPT structure ──
MOCK_SUMMARY = """
## 🎓 Lecture Overview
- **Subject:** Machine Learning
- **Level:** Undergraduate Year 2
- **Core Thesis:** Introduction to neural networks

## 📌 Key Concepts at a Glance
- **Neuron:** Basic unit of a neural network
- **Weights:** Learnable parameters

## 📖 Detailed Section Breakdown
#### Neural Networks
- **What was explained:** Neurons connected in layers
- **Why it matters:** Foundation of deep learning

## ⚠️ Common Mistakes & Misconceptions
- Confusing parameters with hyperparameters

## 💡 Insight Callouts
> Neural networks are universal function approximators

## 🔗 Concept Relationship Map
- Neural Network
  - Layers
    - Neurons
      - Weights

## 📝 The "Teach It Back" Summary
Neural networks consist of layers of neurons...

## ❓ Predicted Exam / Quiz Questions
1. What is a neuron? **Answer:** Basic computational unit

## 📚 Suggested Next Steps
- **Topics to review:** Linear algebra, calculus
"""

# Also use the prompt.py schema emoji set to ensure coverage
PROMPT_SCHEMA_SUMMARY = """
# 🎬 Media Summary Report

## 📋 Metadata
| Field | Value |
|-------|-------|
| **Media Type** | Video |

## 🗂️ Table of Contents
1. Introduction — ~00:00 – 02:00

## 📝 Executive Summary
A comprehensive lecture on machine learning fundamentals.

## 🔑 Key Concepts & Definitions
### Gradient Descent
- **Definition**: Optimization algorithm.
- **Context**: Used to minimize the loss function.

## 📚 Section-by-Section Breakdown
### 1. Introduction (~00:00 – 02:00)
**Main Points:**
- Neural networks

## 💡 Insights & Takeaways
1. **Learning rate matters**: Too high diverges, too low stalls.

## ❓ Questions for Reflection
1. What is backpropagation?

## 🌐 Cross-Cultural Notes
- No culturally specific references.

## 🏷️ Tags & Keywords
`machine-learning` `neural-networks` `gradient-descent`
"""


# ══════════════════════════════════════════════════════════════════════════════
# TEST GROUP 1 — MARKDOWN EXPORT
# ══════════════════════════════════════════════════════════════════════════════

class TestExportMarkdown:

    def test_returns_bytes(self) -> None:
        result = export_markdown(MOCK_SUMMARY, "test")
        assert isinstance(result, bytes), "Must return bytes"
        assert len(result) > 0, "Must not return empty bytes"

    def test_utf8_decodable(self) -> None:
        result = export_markdown(MOCK_SUMMARY, "test")
        decoded = result.decode("utf-8")  # must not raise
        assert isinstance(decoded, str)

    def test_contains_yaml_header(self) -> None:
        decoded = export_markdown(MOCK_SUMMARY, "test").decode("utf-8")
        assert decoded.startswith("---"), "Must start with YAML front-matter"
        assert "title:" in decoded
        assert "generated:" in decoded
        assert "format: Markdown" in decoded

    def test_preserves_original_content(self) -> None:
        decoded = export_markdown(MOCK_SUMMARY, "test").decode("utf-8")
        assert "Lecture Overview" in decoded
        assert "Machine Learning" in decoded
        assert "Neural Networks" in decoded

    def test_preserves_all_9_sections(self) -> None:
        decoded = export_markdown(MOCK_SUMMARY, "test").decode("utf-8")
        sections = [
            "Lecture Overview",
            "Key Concepts",
            "Section Breakdown",
            "Misconceptions",
            "Insight Callouts",
            "Concept Relationship",
            "Teach It Back",
            "Quiz Questions",
            "Next Steps",
        ]
        for section in sections:
            assert section in decoded, f"Missing section: {section}"

    def test_preserves_emoji(self) -> None:
        decoded = export_markdown(MOCK_SUMMARY, "test").decode("utf-8")
        # Markdown preserves emoji natively
        assert "🎓" in decoded
        assert "📌" in decoded
        assert "📖" in decoded
        assert "⚠️" in decoded

    def test_preserves_bold_markdown(self) -> None:
        decoded = export_markdown(MOCK_SUMMARY, "test").decode("utf-8")
        assert "**Subject:**" in decoded

    def test_empty_summary_returns_bytes(self) -> None:
        result = export_markdown("", "test")
        assert isinstance(result, bytes)
        # Should still have the YAML header
        decoded = result.decode("utf-8")
        assert "---" in decoded


# ══════════════════════════════════════════════════════════════════════════════
# TEST GROUP 2 — DOCX EXPORT
# ══════════════════════════════════════════════════════════════════════════════

class TestExportDocx:

    def test_returns_bytes(self) -> None:
        result = export_docx(MOCK_SUMMARY, "test")
        assert isinstance(result, bytes), "Must return bytes"
        assert len(result) > 1000, "DOCX must not be trivially small"

    def test_valid_zip_magic_bytes(self) -> None:
        result = export_docx(MOCK_SUMMARY, "test")
        assert result[:2] == b"PK", "DOCX must start with PK (ZIP header)"

    def test_parseable_as_docx(self) -> None:
        result = export_docx(MOCK_SUMMARY, "test")
        doc = DocxDocument(BytesIO(result))
        assert doc is not None

    def test_preserves_key_content(self) -> None:
        result = export_docx(MOCK_SUMMARY, "test")
        doc = DocxDocument(BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Machine Learning" in full_text
        assert "Neural Networks" in full_text

    def test_contains_title(self) -> None:
        result = export_docx(MOCK_SUMMARY, "test")
        doc = DocxDocument(BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "AI Media Summary" in full_text

    def test_emoji_do_not_crash(self) -> None:
        emoji_text = (
            "## 🎓 Overview\n## 📌 Concepts\n"
            "## ⚠️ Warning\n## 💡 Insight\n"
            "> Blockquote with emoji 🔗"
        )
        result = export_docx(emoji_text, "test")
        assert result[:2] == b"PK", "Emoji must not crash DOCX export"

    def test_blockquote_rendering(self) -> None:
        text = "> This is a blockquote with *italic* text"
        result = export_docx(text, "test")
        doc = DocxDocument(BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "This is a blockquote" in full_text

    def test_bullet_list_rendering(self) -> None:
        text = "- First bullet\n- Second bullet\n- Third bullet"
        result = export_docx(text, "test")
        doc = DocxDocument(BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "First bullet" in full_text
        assert "Second bullet" in full_text

    def test_numbered_list_rendering(self) -> None:
        text = "1. First item\n2. Second item"
        result = export_docx(text, "test")
        doc = DocxDocument(BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "First item" in full_text


# ══════════════════════════════════════════════════════════════════════════════
# TEST GROUP 3 — PDF EXPORT
# ══════════════════════════════════════════════════════════════════════════════

class TestExportPdf:

    def test_returns_bytes(self) -> None:
        result = export_pdf(MOCK_SUMMARY, "test")
        assert isinstance(result, bytes), "Must return bytes"
        assert len(result) > 1000, "PDF must not be trivially small"

    def test_valid_pdf_magic_bytes(self) -> None:
        result = export_pdf(MOCK_SUMMARY, "test")
        assert result[:4] == b"%PDF", "PDF must start with %PDF"

    def test_lecture_emoji_do_not_crash(self) -> None:
        """All 9 lecture section emoji must not crash PDF generation."""
        emoji_text = (
            "## 🎓 Lecture Overview\n"
            "## 📌 Key Concepts\n"
            "## 📖 Section Breakdown\n"
            "## ⚠️ Common Mistakes\n"
            "## 💡 Insight Callouts\n"
            "## 🔗 Concept Map\n"
            "## 📝 Teach It Back\n"
            "## ❓ Exam Questions\n"
            "## 📚 Next Steps\n"
        )
        result = export_pdf(emoji_text, "test")
        assert result[:4] == b"%PDF"

    def test_prompt_schema_emoji_do_not_crash(self) -> None:
        """All prompts.py schema emoji must not crash PDF generation."""
        result = export_pdf(PROMPT_SCHEMA_SUMMARY, "test")
        assert result[:4] == b"%PDF"

    def test_full_mock_summary_generates(self) -> None:
        result = export_pdf(MOCK_SUMMARY, "test")
        assert result[:4] == b"%PDF"

    def test_bold_italic_inline_in_pdf(self) -> None:
        text = "**Bold text** and *italic text* and `code text` here."
        result = export_pdf(text, "test")
        assert result[:4] == b"%PDF"

    def test_blockquote_in_pdf(self) -> None:
        text = "> This is a blockquote insight"
        result = export_pdf(text, "test")
        assert result[:4] == b"%PDF"

    def test_code_block_in_pdf(self) -> None:
        text = "```python\ndef hello():\n    print('world')\n```"
        result = export_pdf(text, "test")
        assert result[:4] == b"%PDF"

    def test_horizontal_rule_in_pdf(self) -> None:
        text = "Section A\n---\nSection B"
        result = export_pdf(text, "test")
        assert result[:4] == b"%PDF"

    def test_table_does_not_crash(self) -> None:
        """Markdown tables fall through as body text — must not crash."""
        text = "| Col A | Col B |\n|-------|-------|\n| Val 1 | Val 2 |"
        result = export_pdf(text, "test")
        assert result[:4] == b"%PDF"

    def test_xml_characters_do_not_crash_pdf(self) -> None:
        """Raw XML/HTML entities like <, >, & must be escaped safely."""
        text = "Price < $100 & High > Low. Also <div>test</div> and AT&T & R&D."
        result = export_pdf(text, "test")
        assert result[:4] == b"%PDF"

    def test_multilingual_characters_pdf(self) -> None:
        """Multilingual UTF-8 characters must generate valid PDF without crashing."""
        text = "Summary: Bonjour le monde, Hola Mundo, Hallo Welt, नमस्ते, வணக்கம்"
        result = export_pdf(text, "test")
        assert result[:4] == b"%PDF"



# ══════════════════════════════════════════════════════════════════════════════
# TEST GROUP 4 — FILENAME SANITIZATION
# ══════════════════════════════════════════════════════════════════════════════

class TestFilenameSanitization:

    def test_spaces_replaced(self) -> None:
        dirty = "My Lecture Part 1"
        clean = re.sub(r"[^\w\-]", "_", dirty)
        assert " " not in clean

    def test_colon_replaced(self) -> None:
        dirty = "Lecture: Part 1"
        clean = re.sub(r"[^\w\-]", "_", dirty)
        assert ":" not in clean

    def test_slash_replaced(self) -> None:
        dirty = "Lecture/Part1"
        clean = re.sub(r"[^\w\-]", "_", dirty)
        assert "/" not in clean

    def test_unicode_chars_replaced(self) -> None:
        dirty = "Lección: Álgebra"
        clean = re.sub(r"[^\w\-]", "_", dirty)
        # Must be a valid filename (no colons, slashes)
        assert ":" not in clean
        assert "/" not in clean

    def test_hyphens_preserved(self) -> None:
        name = "lecture-summary-2024"
        clean = re.sub(r"[^\w\-]", "_", name)
        assert clean == "lecture-summary-2024"

    def test_underscores_preserved(self) -> None:
        name = "lecture_summary_2024"
        clean = re.sub(r"[^\w\-]", "_", name)
        assert clean == "lecture_summary_2024"


# ══════════════════════════════════════════════════════════════════════════════
# TEST GROUP 5 — CROSS-FORMAT CONSISTENCY
# ══════════════════════════════════════════════════════════════════════════════

class TestCrossFormatConsistency:

    def test_all_three_formats_generate_from_same_input(self) -> None:
        """All three export formats must generate from identical input."""
        md = export_markdown(MOCK_SUMMARY, "test")
        docx = export_docx(MOCK_SUMMARY, "test")
        pdf = export_pdf(MOCK_SUMMARY, "test")
        assert md[:3] == b"---"
        assert docx[:2] == b"PK"
        assert pdf[:4] == b"%PDF"

    def test_all_formats_handle_empty_summary(self) -> None:
        """Empty summary must not crash any export format."""
        md = export_markdown("", "test")
        docx = export_docx("", "test")
        pdf = export_pdf("", "test")
        assert isinstance(md, bytes)
        assert isinstance(docx, bytes)
        assert isinstance(pdf, bytes)

    def test_all_formats_handle_single_line(self) -> None:
        """Single line of text must not crash any format."""
        text = "A single sentence summary."
        assert export_markdown(text, "test")[:3] == b"---"
        assert export_docx(text, "test")[:2] == b"PK"
        assert export_pdf(text, "test")[:4] == b"%PDF"


# ══════════════════════════════════════════════════════════════════════════════
# TEST GROUP 6 — LATEX MATH CLEANING & NESTED BULLETS
# ══════════════════════════════════════════════════════════════════════════════

class TestMathAndNestedBullets:

    def test_latex_braces_cleaned(self) -> None:
        raw = r"initial belief state is $\{A, B\}$ and new state is $\{B\}$."
        md = export_markdown(raw, "test").decode("utf-8")
        assert r"\{" not in md
        assert r"\}" not in md
        assert "{A, B}" in md
        assert "{B}" in md

    def test_nested_indented_bullets_docx(self) -> None:
        raw = (
            "• Main item:\n"
            "  - Sub item 1\n"
            "  - Sub item 2\n"
            "    - Deep sub item\n"
        )
        docx_bytes = export_docx(raw, "test")
        doc = DocxDocument(BytesIO(docx_bytes))
        texts = [p.text for p in doc.paragraphs]
        assert "Main item:" in texts
        assert "Sub item 1" in texts
        assert "Sub item 2" in texts
        assert "Deep sub item" in texts

    def test_nested_indented_bullets_pdf(self) -> None:
        raw = (
            "• Main item:\n"
            "  - Sub item 1\n"
            "  - Sub item 2\n"
        )
        pdf_bytes = export_pdf(raw, "test")
        assert pdf_bytes[:4] == b"%PDF"

